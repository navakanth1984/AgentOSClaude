"""
creative_exporter.py — Export Screenplays and Novels from Agent OS to DOCX, PDF, HTML, and MD
Handles Hollywood standard screenplay indentation/formatting and novel styles.
Uses python-docx for Word files and Playwright for high-fidelity PDF files.
"""

import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Playwright optional import
try:
    from playwright.sync_api import sync_playwright
    HAS_PLAYWRIGHT = True
except ImportError:
    HAS_PLAYWRIGHT = False


# ── Screenplay Parser ─────────────────────────────────────────────────────────

def parse_markdown_screenplay(md_text: str) -> list[dict]:
    """
    Parses a markdown screenplay into structured blocks.
    Identifies: scene_header, action, character, parenthetical, dialogue.
    """
    lines = md_text.splitlines()
    blocks = []
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
            
        # 1. Scene Header
        # e.g. **INT. ALLEYWAY - NIGHT** or EXT. ROOFTOP
        if (line.startswith("**") and ("INT." in line or "EXT." in line)) or \
           (line.upper().startswith("INT. ") or line.upper().startswith("EXT. ") or \
            line.upper().startswith("INT/EXT. ") or line.upper().startswith("EXT/INT. ")):
            clean_header = line.replace("**", "").strip().upper()
            blocks.append({"type": "scene_header", "text": clean_header})
            i += 1
            continue
            
        # 2. Character + Dialogue (Colon format)
        # e.g. KAMRAN: We need to move.
        match_colon = re.match(r"^([A-Z0-9\s\-\(\)\.]+)\s*:\s*(.*)$", line)
        if match_colon:
            char_name = match_colon.group(1).strip().upper()
            dialogue_text = match_colon.group(2).strip()
            
            blocks.append({"type": "character", "text": char_name})
            
            # Check if dialogue starts with a parenthetical
            if dialogue_text.startswith("("):
                p_match = re.match(r"^(\([^\)]+\))\s*(.*)$", dialogue_text)
                if p_match:
                    blocks.append({"type": "parenthetical", "text": p_match.group(1).strip()})
                    dialogue_text = p_match.group(2).strip()
            
            if dialogue_text:
                blocks.append({"type": "dialogue", "text": dialogue_text})
            i += 1
            continue

        # 3. Standard Character Name
        # Check if line is uppercase, short, and the next line is a parenthetical or dialogue
        if line.isupper() and len(line) < 30 and not any(x in line for x in [".", ",", "!", "?"]):
            next_line = lines[i+1].strip() if i+1 < len(lines) else ""
            # If the next line is a parenthetical or simple dialogue block
            if next_line.startswith("(") or (next_line and not next_line.isupper()):
                blocks.append({"type": "character", "text": line})
                i += 1
                continue

        # 4. Parenthetical
        # e.g. (whispering)
        if line.startswith("(") and line.endswith(")"):
            blocks.append({"type": "parenthetical", "text": line})
            i += 1
            continue

        # 5. Default: Action Line
        blocks.append({"type": "action", "text": line})
        i += 1
        
    return blocks


# ── HTML Exporter ─────────────────────────────────────────────────────────────

def export_to_html(blocks: list[dict], is_screenplay: bool = True) -> str:
    """Generate screenplay formatted or novel formatted HTML."""
    if is_screenplay:
        css = """
        body {
            background-color: #f7f7f7;
            font-family: "Courier Prime", "Courier New", Courier, monospace;
            font-size: 12pt;
            line-height: 1.15;
            color: #000;
            margin: 0;
            padding: 0;
        }
        .page {
            background-color: #fff;
            width: 8.5in;
            min-height: 11in;
            margin: 0.5in auto;
            padding: 1.0in 1.0in 1.0in 1.5in; /* Standard Hollywood margins */
            box-sizing: border-box;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        .scene_header {
            text-transform: uppercase;
            font-weight: bold;
            margin-top: 24px;
            margin-bottom: 12px;
            text-align: left;
            page-break-after: avoid;
        }
        .action {
            margin-top: 12px;
            margin-bottom: 12px;
            text-align: left;
        }
        /* Dialogue block is a narrow column CENTERED on the page (equal L/R
           indents), with the words inside left-aligned so wrapping reads
           naturally — the industry-standard screenplay look. The CHARACTER
           cue is centered directly above that column. */
        .character {
            text-transform: uppercase;
            margin-left: 1.5in;
            margin-right: 1.5in;
            text-align: center;
            margin-top: 12px;
            margin-bottom: 0px;
            page-break-after: avoid;
        }
        .parenthetical {
            margin-left: 2.0in;
            margin-right: 2.0in;
            text-align: left;
            margin-top: 0px;
            margin-bottom: 0px;
            page-break-after: avoid;
        }
        .dialogue {
            margin-left: 1.5in;
            margin-right: 1.5in;
            text-align: left;
            margin-top: 0px;
            margin-bottom: 12px;
        }
        @media print {
            body {
                background-color: #fff;
            }
            .page {
                width: auto;
                min-height: auto;
                margin: 0;
                padding: 1.0in 1.0in 1.0in 1.5in;
                box-shadow: none;
                page-break-after: always;
            }
        }
        """
        body_content = ""
        for b in blocks:
            cls = b["type"]
            text = b["text"]
            body_content += f'<div class="{cls}">{text}</div>\n'
    else:
        # Novel Formatting CSS
        css = """
        body {
            background-color: #f7f7f7;
            font-family: "Georgia", serif;
            font-size: 11.5pt;
            line-height: 1.6;
            color: #111;
            margin: 0;
            padding: 0;
        }
        .page {
            background-color: #fff;
            width: 6.0in;
            min-height: 9.0in;
            margin: 0.5in auto;
            padding: 1.0in;
            box-sizing: border-box;
            box-shadow: 0 0 10px rgba(0,0,0,0.1);
        }
        h1, h2, h3 {
            font-family: "Palatino", "Georgia", serif;
            text-align: center;
            margin-top: 30px;
            margin-bottom: 20px;
        }
        p {
            text-indent: 0.5in;
            margin: 0 0 10px 0;
            text-align: justify;
        }
        p.first-para {
            text-indent: 0;
        }
        @media print {
            body {
                background-color: #fff;
            }
            .page {
                width: auto;
                min-height: auto;
                margin: 0;
                padding: 1.0in;
                box-shadow: none;
                page-break-after: always;
            }
        }
        """
        # Parse novel text (paragraphs)
        body_content = ""
        first = True
        for line in blocks:
            text = line["text"]
            if line["type"] == "scene_header" or text.startswith("#"):
                clean_h = text.replace("#", "").strip()
                body_content += f'<h2>{clean_h}</h2>\n'
                first = True
            else:
                cls = "first-para" if first else ""
                body_content += f'<p class="{cls}">{text}</p>\n'
                first = False

    html_template = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>Exported Output</title>
    <style>{css}</style>
</head>
<body>
    <div class="page">
        {body_content}
    </div>
</body>
</html>
"""
    return html_template


# ── PDF Exporter (Playwright) ──────────────────────────────────────────────────

def export_to_pdf(html_content: str, output_path: str) -> bool:
    """Uses Playwright to print the HTML content to PDF with exact formatting."""
    if not HAS_PLAYWRIGHT:
        print("[Exporter] Error: Playwright package not installed.")
        return False
        
    temp_html = Path(output_path).parent / "temp_print.html"
    try:
        temp_html.write_text(html_content, encoding="utf-8")
        
        with sync_playwright() as p:
            # Launch headless chromium
            browser = p.chromium.launch()
            page = browser.new_page()
            # Navigate to local temp html
            page.goto(temp_html.absolute().as_uri())
            
            # Print page to PDF with standard margins matching HTML styles
            page.pdf(
                path=output_path,
                format="Letter",
                print_background=True,
                margin={"top": "0in", "bottom": "0in", "left": "0in", "right": "0in"} # Margins handled in CSS
            )
            browser.close()
            
        print(f"[Exporter] Successfully printed PDF to {output_path}")
        return True
    except Exception as e:
        print(f"[Exporter] PDF Generation error: {e}")
        return False
    finally:
        if temp_html.exists():
            temp_html.unlink()


# ── DOCX Exporter (python-docx) ───────────────────────────────────────────────

def export_to_docx(blocks: list[dict], output_path: str, is_screenplay: bool = True) -> bool:
    """Generate Word file preserving exact screenplay indentation or novel styles."""
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    if is_screenplay:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.left_margin = Inches(1.5)  # Screenplay left margin (binding)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        
        # Configure styles
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(12)
        
        for b in blocks:
            text = b["text"]
            b_type = b["type"]
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.15
            
            if b_type == "scene_header":
                p_format.left_indent = Inches(0.0)
                p_format.space_before = Pt(18)
                p_format.space_after = Pt(12)
                run = p.add_run(text.upper())
                run.bold = True
            elif b_type == "action":
                p_format.left_indent = Inches(0.0)
                p_format.space_before = Pt(6)
                p_format.space_after = Pt(6)
                p.add_run(text)
            elif b_type == "character":
                # Centered directly above the centered dialogue column: equal
                # L/R indents + CENTER alignment place the cue over the column.
                p_format.left_indent = Inches(1.5)
                p_format.right_indent = Inches(1.5)
                p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_format.space_before = Pt(12)
                p_format.space_after = Pt(0)
                p.add_run(text.upper())
            elif b_type == "parenthetical":
                # Sits inside the dialogue column, slightly narrower, flush-left.
                p_format.left_indent = Inches(2.0)
                p_format.right_indent = Inches(2.0)
                p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(0)
                p.add_run(text)
            elif b_type == "dialogue":
                # Narrow column CENTERED on the page (equal 1.5in L/R indents),
                # with the text left-aligned so wrapped lines stay flush-left.
                p_format.left_indent = Inches(1.5)
                p_format.right_indent = Inches(1.5)
                p_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p_format.space_before = Pt(0)
                p_format.space_after = Pt(6)
                p.add_run(text)
    else:
        # Novel Page setup
        section.page_width = Inches(6.0)
        section.page_height = Inches(9.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Georgia'
        font.size = Pt(11.5)
        
        first = True
        for b in blocks:
            text = b["text"]
            b_type = b["type"]
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.line_spacing = 1.3
            p_format.space_after = Pt(6)
            
            if b_type == "scene_header" or text.startswith("#"):
                p_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_format.space_before = Pt(24)
                p_format.space_after = Pt(12)
                run = p.add_run(text.replace("#", "").strip())
                run.bold = True
                run.font.size = Pt(16)
                first = True
            else:
                if not first:
                    p_format.first_line_indent = Inches(0.5)
                p.add_run(text)
                first = False
                
    try:
        doc.save(output_path)
        print(f"[Exporter] Successfully saved DOCX to {output_path}")
        return True
    except Exception as e:
        print(f"[Exporter] DOCX Generation error: {e}")
        return False


# ── Main Export Interface ────────────────────────────────────────────────────

def export_document(md_content: str, base_path: str, is_screenplay: bool = True) -> dict:
    """
    Exports a screenplay or novel to MD, HTML, DOCX, and PDF formats.
    Returns a dict with paths to all generated formats.
    """
    base_p = Path(base_path)
    base_p.parent.mkdir(parents=True, exist_ok=True)
    
    # Parse blocks
    blocks = parse_markdown_screenplay(md_content)
    
    results = {}
    
    # 1. MD format
    md_path = base_p.with_suffix(".md")
    md_path.write_text(md_content, encoding="utf-8")
    results["md"] = str(md_path)
    
    # 2. HTML format
    html_content = export_to_html(blocks, is_screenplay)
    html_path = base_p.with_suffix(".html")
    html_path.write_text(html_content, encoding="utf-8")
    results["html"] = str(html_path)
    
    # 3. DOCX format
    docx_path = base_p.with_suffix(".docx")
    export_to_docx(blocks, str(docx_path), is_screenplay)
    results["docx"] = str(docx_path)
    
    # 4. PDF format
    if HAS_PLAYWRIGHT:
        pdf_path = base_p.with_suffix(".pdf")
        export_to_pdf(html_content, str(pdf_path))
        results["pdf"] = str(pdf_path)
    else:
        results["pdf"] = "Not generated (Playwright missing)"
        
    return results
